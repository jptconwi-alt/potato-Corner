import os
import uuid
from flask import render_template, request, redirect, url_for, flash, jsonify, session
from flask_login import login_required, current_user, logout_user
from datetime import datetime
from werkzeug.utils import secure_filename
from models import ph_now
from models import db, User, Product, Order, OrderItem, CartItem, OrderRating
from controllers import AuthController, ProductController, CartController, OrderController
from auth_decorator import admin_required

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
UPLOAD_FOLDER = os.path.join('static', 'images', 'uploads')


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_session_id():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        session.modified = True   # force Flask to save the cookie
    session.permanent = True      # survive browser restarts (uses PERMANENT_SESSION_LIFETIME)
    return session['session_id']


def register_routes(app):
    try:
        os.makedirs(os.path.join(app.root_path, UPLOAD_FOLDER), exist_ok=True)
    except OSError:
        pass  # Read-only filesystem (Vercel) — file uploads via disk unavailable

    # ─────────────────────────────────────────
    # PUBLIC ROUTES
    # ─────────────────────────────────────────

    @app.route('/')
    def index():
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        products = Product.query.order_by(Product.flavor, Product.size).all()
        flavors = ProductController.get_flavors()
        return render_template('index.html', products=products, flavors=flavors)

    # ─────────────────────────────────────────
    # AUTH ROUTES
    # ─────────────────────────────────────────



    # ── AJAX Login (used by inline auth modal) ──────────
    @app.route('/auth/ajax-login', methods=['POST'])
    def ajax_login():
        data     = request.get_json(force=True) or {}
        username = data.get('username', '').strip()
        password = data.get('password', '')
        if not username or not password:
            return jsonify({'success': False, 'message': 'Username and password are required.'})
        success, result = AuthController.login_user(username, password, False)
        if success:
            session['user_id'] = result.id
            sid = get_session_id()
            CartController.merge_carts(sid, result.id)
            return jsonify({'success': True, 'full_name': result.full_name})
        return jsonify({'success': False, 'message': str(result)})

    # ── AJAX Register (used by inline auth modal) ───────
    @app.route('/auth/ajax-register', methods=['POST'])
    def ajax_register():
        data      = request.get_json(force=True) or {}
        firstname = data.get('firstname', '').strip()
        lastname  = data.get('lastname', '').strip()
        username  = data.get('username', '').strip()
        email     = data.get('email', '').strip()
        phone     = data.get('phone', '').strip()
        password  = data.get('password', '')
        full_name = f'{firstname} {lastname}'.strip()
        if not all([firstname, lastname, username, email, password]):
            return jsonify({'success': False, 'message': 'Please fill in all required fields.'})
        if len(password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters.'})
        success, result = AuthController.register_user(username, email, password, full_name, phone)
        if success:
            # Auto-login after registration
            success2, result2 = AuthController.login_user(username, password, False)
            if success2:
                session['user_id'] = result2.id
            return jsonify({'success': True, 'full_name': full_name})
        return jsonify({'success': False, 'message': str(result)})

    # ── Login page ───────────────────────────────────────
    @app.route('/login', methods=['GET', 'POST'])
    def login():
        if current_user.is_authenticated:
            return redirect(url_for('index'))
        if request.method == 'POST':
            username = request.form.get('username', '').strip()
            password = request.form.get('password', '')
            remember = bool(request.form.get('remember'))
            if not username or not password:
                flash('Username and password are required.', 'danger')
                return render_template('login.html')
            success, result = AuthController.login_user(username, password, remember)
            if success:
                session['user_id'] = result.id
                sid = get_session_id()
                CartController.merge_carts(sid, result.id)
                return redirect(url_for('index'))
            flash(str(result), 'danger')
        return render_template('login.html')

    # ── Register page ────────────────────────────────────
    @app.route('/register', methods=['GET', 'POST'])
    def register():
        if current_user.is_authenticated:
            return redirect(url_for('index'))
        if request.method == 'POST':
            username = request.form.get('username', '').strip()
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '')
            confirm = request.form.get('confirm_password', '')
            full_name = request.form.get('full_name', '').strip()
            phone = request.form.get('phone', '').strip()
            street = request.form.get('street', '').strip()
            barangay = request.form.get('barangay', '').strip()
            city = request.form.get('city', '').strip()
            province = request.form.get('province', '').strip()
            zipcode = request.form.get('zipcode', '').strip()
            if password != confirm:
                flash('Passwords do not match.', 'danger')
                return render_template('register.html')
            if len(password) < 6:
                flash('Password must be at least 6 characters.', 'danger')
                return render_template('register.html')
            success, result = AuthController.register_user(username, email, password, full_name, phone)
            if success:
                result.street = street
                result.barangay = barangay
                result.city = city
                result.province = province
                result.zipcode = zipcode
                result.profile_complete = True
                db.session.commit()
                success2, result2 = AuthController.login_user(username, password, False)
                if success2:
                    session['user_id'] = result2.id
                    sid = get_session_id()
                    CartController.merge_carts(sid, result2.id)
                flash('Account created! Welcome to Potato Corner 🍟', 'success')
                return redirect(url_for('index'))
            flash(str(result), 'danger')
        return render_template('register.html')

    @app.route('/logout')
    def logout():
        # Do NOT re-assign cart items to the session on logout.
        # Doing so causes a critical bug: after a user places an order,
        # clear_selected_items() removes the ordered cart rows by ID.
        # If logout then re-attaches the (already-deleted) rows — or any
        # leftover rows — to the session, those items reappear in the cart
        # on the next page refresh even though the order was placed.
        # Cart items stay linked to user_id; they will be restored when the
        # user logs back in via merge_carts().
        logout_user()          # Flask-Login: clears user, queues remember cookie deletion
        session.pop('user_id', None)   # remove only our custom key, NOT session.clear()
        session.modified = True
        flash('You have been logged out.', 'info')
        response = redirect(url_for('login'))
        # Explicitly expire the remember_token cookie that Flask-Login sets
        response.delete_cookie('remember_token')
        return response

    @app.route('/complete-profile', methods=['GET', 'POST'])
    @login_required
    def complete_profile():
        if request.method == 'POST':
            current_user.full_name = request.form.get('full_name', current_user.full_name).strip()
            current_user.phone = request.form.get('phone', '').strip()
            current_user.street = request.form.get('street', '').strip()
            current_user.barangay = request.form.get('barangay', '').strip()
            current_user.city = request.form.get('city', '').strip()
            current_user.province = request.form.get('province', '').strip()
            current_user.zipcode = request.form.get('zipcode', '').strip()
            current_user.profile_complete = True
            db.session.commit()
            flash('Profile complete! Start ordering.', 'success')
            return redirect(url_for('index'))
        return render_template('complete_profile.html')

    # ─────────────────────────────────────────
    # PROFILE & ORDERS
    # ─────────────────────────────────────────

    @app.route('/profile', methods=['GET', 'POST'])
    @login_required
    def profile():
        orders = OrderController.get_user_orders(current_user.id)
        if request.method == 'POST':
            current_user.full_name = request.form.get('full_name', current_user.full_name).strip()
            current_user.phone = request.form.get('phone', current_user.phone or '').strip()
            current_user.street = request.form.get('street', current_user.street or '').strip()
            current_user.barangay = request.form.get('barangay', current_user.barangay or '').strip()
            current_user.city = request.form.get('city', current_user.city or '').strip()
            new_email = request.form.get('email', '').strip()
            if new_email and new_email != current_user.email:
                if User.query.filter_by(email=new_email).first():
                    flash('Email already in use', 'danger')
                    return render_template('profile.html', orders=orders)
                current_user.email = new_email
            db.session.commit()
            flash('Profile updated!', 'success')
        return render_template('profile.html', orders=orders)

    @app.route('/orders')
    @login_required
    def orders():
        user_orders = OrderController.get_user_orders(current_user.id)
        return render_template('my_orders.html', orders=user_orders)

    @app.route('/my-orders')
    @login_required
    def my_orders():
        return redirect(url_for('orders'))

    @app.route('/track-order')
    def track_order():
        order_number = request.args.get('order_number', '').strip()
        order = None
        if order_number:
            order = OrderController.get_order(order_number)
            if not order:
                flash('Order not found. Please check the order number.', 'danger')
        return render_template('track_order.html', order=order, order_number=order_number)

    # ─────────────────────────────────────────
    # CART ROUTES
    # ─────────────────────────────────────────

    @app.route('/cart')
    def cart():
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        blacklist = session.get('ordered_item_ids', [])
        # Expire the blacklist after 5 minutes regardless — prevents it from
        # hiding legitimately re-added items if Turso lag is unusually long.
        import time as _time
        bl_ts = session.get('ordered_item_ids_ts', 0)
        if blacklist and (_time.time() - bl_ts) > 300:
            blacklist = []
            session.pop('ordered_item_ids', None)
            session.pop('ordered_item_ids_ts', None)
            session.modified = True

        # Fetch ALL rows first (without exclusion) so we can detect whether
        # Turso is still returning stale rows for the ordered IDs.
        all_items = CartController.get_cart_items(sid, uid)
        if blacklist:
            bl_set = set(blacklist)
            # Only clear the blacklist once Turso is no longer returning ANY of
            # the ordered rows — i.e. the delete has fully propagated.
            still_stale = any(i.id in bl_set for i in all_items)
            if not still_stale:
                session.pop('ordered_item_ids', None)
                session.pop('ordered_item_ids_ts', None)
                session.modified = True
            cart_items = [i for i in all_items if i.id not in bl_set]
        else:
            cart_items = all_items
        subtotal = sum(i.product.price * i.quantity for i in cart_items)
        delivery_fee = 0 if subtotal >= 500 else 50
        total = subtotal + delivery_fee
        return render_template('cart.html', cart_items=cart_items, subtotal=subtotal, delivery_fee=delivery_fee, total=total)

    @app.route('/cart/add', methods=['POST'])
    def cart_add():
        data = request.get_json()
        product_id = data.get('product_id')
        quantity = int(data.get('quantity', 1))
        product = Product.query.get(product_id)
        if not product or not product.is_available:
            return jsonify({'success': False, 'message': 'Product not available'})
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        CartController.add_to_cart(sid, product_id, uid, quantity)
        return jsonify({'success': True, 'message': 'Added to cart'})

    @app.route('/cart/update', methods=['POST'])
    def cart_update():
        data = request.get_json()
        item_id = data.get('item_id')
        quantity = int(data.get('quantity', 1))
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        ok = CartController.update_quantity(sid, item_id, quantity, uid)
        if not ok:
            # Ownership check failed — item may belong to this user but the
            # session_id stored on the row is stale (e.g. added before login).
            # Retry with a pure user_id-only lookup to handle that edge case.
            if uid:
                from models import CartItem as _CI
                item = _CI.query.filter_by(id=item_id, user_id=uid).first()
                if item:
                    if quantity <= 0:
                        db.session.delete(item)
                    else:
                        item.quantity = quantity
                    db.session.commit()
                    ok = True
        return jsonify({'success': bool(ok)})

    @app.route('/cart/remove', methods=['POST'])
    def cart_remove():
        data = request.get_json()
        item_id = data.get('item_id')
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        ok = CartController.remove_from_cart(sid, item_id, uid)
        return jsonify({'success': bool(ok)})

    @app.route('/api/cart/count')
    def cart_count():
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        # If the user is logged in, purge any leftover session-only rows that
        # were not cleaned up by merge_carts (e.g. from an old anonymous session
        # stored in the same browser). These are the root cause of badge > actual.
        if uid:
            from models import CartItem as _CI
            stale = _CI.query.filter(
                _CI.session_id == sid,
                _CI.user_id == None  # noqa: E711
            ).all()
            if stale:
                for s in stale:
                    db.session.delete(s)
                db.session.commit()
        blacklist = session.get('ordered_item_ids', [])
        cart_items = CartController.get_cart_items(sid, uid, exclude_ids=blacklist)
        total = sum(i.quantity for i in cart_items)
        return jsonify({'count': total})

    @app.route('/api/cart/items')
    def api_cart_items():
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None
        blacklist = session.get('ordered_item_ids', [])
        cart_items = CartController.get_cart_items(sid, uid, exclude_ids=blacklist)
        items = []
        for i in cart_items:
            items.append({
                'id': i.id,
                'product_id': i.product_id,
                'product_name': i.product.name,
                'product_price': float(i.product.price),
                'quantity': i.quantity,
                'subtotal': float(i.product.price * i.quantity),
            })
        subtotal = sum(i['subtotal'] for i in items)
        delivery_fee = 0 if subtotal >= 500 else 50
        return jsonify({
            'items': items,
            'count': sum(i['quantity'] for i in items),
            'subtotal': subtotal,
            'delivery_fee': delivery_fee,
            'total': subtotal + delivery_fee,
        })

    @app.route('/api/orders/status')
    def api_orders_status():
        if not current_user.is_authenticated:
            return jsonify({'orders': []})
        user_orders = OrderController.get_user_orders(current_user.id)
        orders = []
        for o in user_orders:
            orders.append({
                'id': o.id,
                'order_number': o.order_number,
                'status': o.status,
                'total': float(o.total_amount),
                'created_at': o.order_date.isoformat() if o.order_date else None,
                'items_count': sum(i.quantity for i in o.items),
            })
        return jsonify({'orders': orders})

    @app.route('/api/admin/orders/poll')
    @admin_required
    def api_admin_orders_poll():
        """Admin long-poll endpoint — returns current status of all orders.
        The admin orders page calls this every 4 s to stay in sync across
        browser tabs without Socket.IO."""
        orders = Order.query.order_by(Order.order_date.desc()).all()
        return jsonify({'orders': [
            {'id': o.id, 'order_number': o.order_number, 'status': o.status,
             'payment_status': o.payment_status}
            for o in orders
        ]})

    # ─────────────────────────────────────────
    # CHECKOUT & ORDERS
    # ─────────────────────────────────────────

    @app.route('/checkout', methods=['GET', 'POST'])
    def checkout():
        sid = get_session_id()
        uid = current_user.id if current_user.is_authenticated else None

        blacklist = session.get('ordered_item_ids', [])
        all_cart_items = CartController.get_cart_items(sid, uid, exclude_ids=blacklist)
        if not all_cart_items:
            flash('Your cart is empty', 'warning')
            return redirect(url_for('cart'))

        if request.method == 'GET':
            # Read selected item IDs from query string (sent by cart page)
            selected_ids_raw = request.args.getlist('items')
            selected_ids = [int(x) for x in selected_ids_raw if x.isdigit()]

            if selected_ids:
                cart_items = [i for i in all_cart_items if i.id in selected_ids]
            else:
                cart_items = all_cart_items

            if not cart_items:
                flash('No items selected. Please select items to order.', 'warning')
                return redirect(url_for('cart'))

            # Persist the selected IDs in the session so the POST handler
            # can reliably read them even if the query string is lost.
            session['checkout_item_ids'] = [i.id for i in cart_items]
            session.modified = True

        else:  # POST
            # Prefer IDs from query string, then session fallback
            selected_ids_raw = request.args.getlist('items') or request.form.getlist('items')
            selected_ids = [int(x) for x in selected_ids_raw if x.isdigit()]

            if not selected_ids:
                # Fallback: use what was stored in session during GET
                selected_ids = session.get('checkout_item_ids', [])

            if selected_ids:
                cart_items = [i for i in all_cart_items if i.id in selected_ids]
            else:
                cart_items = all_cart_items

            if not cart_items:
                flash('Your cart is empty or items were already ordered.', 'warning')
                return redirect(url_for('cart'))

        subtotal = sum(i.product.price * i.quantity for i in cart_items)
        delivery_fee = 0 if subtotal >= 500 else 50
        total = subtotal + delivery_fee

        # Carry selected IDs into the POST form so the POST handler knows which to clear
        items_param = '&'.join(f'items={i.id}' for i in cart_items)

        if request.method == 'POST':
            customer_data = {
                'name': request.form.get('name', '').strip(),
                'email': request.form.get('email', '').strip(),
                'phone': request.form.get('phone', '').strip(),
                'address': request.form.get('address', '').strip(),
                'payment_method': request.form.get('payment_method', 'Cash on Delivery'),
                'delivery_lat': float(request.form.get('delivery_lat')) if request.form.get('delivery_lat') else None,
                'delivery_lng': float(request.form.get('delivery_lng')) if request.form.get('delivery_lng') else None,
            }
            if not all([customer_data['name'], customer_data['email'], customer_data['phone'], customer_data['address']]):
                flash('Please fill in all required fields', 'danger')
                return render_template('checkout.html', cart_items=cart_items, subtotal=subtotal,
                                       delivery_fee=delivery_fee, total=total, items_param=items_param)

            order = OrderController.create_order(sid, customer_data, cart_items, uid, delivery_fee)

            # Clear ONLY the ordered items (not unselected items left in cart).
            # clear_selected_items: soft-deletes (is_ordered=True), hard-deletes,
            # and calls db.session.expire_all() to flush any stale ORM cache.
            ordered_ids = [i.id for i in cart_items]
            CartController.clear_selected_items(sid, uid, ordered_ids)

            # Session-level blacklist as an extra safety net for replication lag.
            import time as _time
            existing_blacklist = session.get('ordered_item_ids', [])
            session['ordered_item_ids'] = list(set(existing_blacklist + ordered_ids))
            session['ordered_item_ids_ts'] = _time.time()
            session.pop('checkout_item_ids', None)
            session.modified = True

            return redirect(url_for('order_confirmation', order_number=order.order_number))

        return render_template('checkout.html', cart_items=cart_items, subtotal=subtotal,
                               delivery_fee=delivery_fee, total=total, items_param=items_param)

    @app.route('/order/<int:order_id>/cancel', methods=['POST'])
    @login_required
    def cancel_order(order_id):
        order = Order.query.get_or_404(order_id)
        # Only the owner can cancel their own order
        if order.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        # Only allow cancellation of Pending or Confirmed orders
        if order.status not in ['Pending', 'Confirmed']:
            return jsonify({'success': False, 'message': f'Cannot cancel an order that is already "{order.status}".'})
        order.status = 'Cancelled'
        db.session.commit()
        return jsonify({'success': True, 'message': 'Order cancelled successfully.'})

    @app.route('/order/confirmation/<order_number>')
    def order_confirmation(order_number):
        order = OrderController.get_order(order_number)
        if not order:
            flash('Order not found', 'danger')
            return redirect(url_for('index'))
        return render_template('order_confirmation.html', order=order)

    # ─────────────────────────────────────────
    # ADMIN ROUTES
    # ─────────────────────────────────────────

    # ── Emergency admin reset (visit URL to force-reset credentials) ─────────
    @app.route('/admin/reset-credentials')
    def admin_reset_credentials():
        """Force-reset admin password. Remove this route after use."""
        try:
            from init_db import init_database
            init_database()
            return '<h2>✅ Admin credentials reset!</h2><p>Username: <b>admin</b> &nbsp; Password: <b>Admin@1234</b></p><a href="/admin/login">Go to Admin Login</a>'
        except Exception as e:
            return f'<h2>❌ Error: {e}</h2>', 500

    @app.route('/admin/login', methods=['GET', 'POST'])
    def admin_login():
        # Already logged in as admin — go straight to dashboard
        if current_user.is_authenticated and current_user.is_admin:
            return redirect(url_for('admin_dashboard'))
        if request.method == 'POST':
            username = request.form.get('username', '').strip()
            password = request.form.get('password', '')
            success, result = AuthController.login_user(username, password, False)
            if success and result.is_admin:
                session['user_id'] = result.id
                return redirect(url_for('admin_dashboard'))
            else:
                # Wrong credentials OR not an admin
                if success and not result.is_admin:
                    from flask_login import logout_user as _logout
                    _logout()  # log them back out immediately
                flash('Invalid admin credentials. Access denied.', 'danger')
        return render_template('admin_login.html')

    @app.route('/admin')
    @admin_required
    def admin_dashboard():
        all_orders = Order.query.all()
        stats = {
            'total_orders': len(all_orders),
            'pending_orders': sum(1 for o in all_orders if o.status == 'Pending'),
            'delivered_orders': sum(1 for o in all_orders if o.status == 'Delivered'),
            'out_for_delivery': sum(1 for o in all_orders if o.status == 'Out for Delivery'),
            'total_revenue': sum(o.total_amount for o in all_orders if o.status == 'Delivered'),
            'total_users': User.query.count(),
        }
        recent_orders = Order.query.order_by(Order.order_date.desc()).limit(8).all()
        pending_delivery = Order.query.filter(
            Order.status.in_(['Pending', 'Preparing', 'Out for Delivery'])
        ).order_by(Order.order_date.asc()).limit(5).all()
        pending_count = stats['pending_orders']
        return render_template('admin/dashboard.html',
                               stats=stats,
                               recent_orders=recent_orders,
                               pending_delivery=pending_delivery,
                               pending_count=pending_count,
                               now=ph_now())

    @app.route('/admin/orders')
    @admin_required
    def admin_orders():
        all_orders = Order.query.order_by(Order.order_date.desc()).all()
        return render_template('admin/orders.html', orders=all_orders)

    @app.route('/admin/users')
    @admin_required
    def admin_users():
        all_users = User.query.order_by(User.created_at.desc()).all()
        return render_template('admin/users.html', users=all_users)

    @app.route('/admin/user/<int:user_id>/toggle-active', methods=['POST'])
    @admin_required
    def admin_toggle_user_active(user_id):
        user = User.query.get_or_404(user_id)
        if user.is_admin:
            return jsonify({'success': False, 'message': 'Cannot disable an admin account.'})
        user.is_active = not user.is_active
        db.session.commit()
        state = 'enabled' if user.is_active else 'disabled'
        return jsonify({'success': True, 'is_active': user.is_active, 'message': f'User {state} successfully.'})

    @app.route('/admin/products')
    @admin_required
    def admin_products():
        all_products = Product.query.order_by(Product.flavor, Product.size).all()
        return render_template('admin/products.html', products=all_products)

    # ── Admin AJAX: Update order status ─────

    @app.route('/admin/order/<int:order_id>/status', methods=['POST'])
    @admin_required
    def admin_update_order_status(order_id):
        data = request.get_json()
        new_status = data.get('status')
        valid = ['Pending', 'Preparing', 'Out for Delivery', 'Delivered', 'Cancelled', 'Confirmed']
        if new_status not in valid:
            return jsonify({'success': False, 'message': 'Invalid status'})
        success = OrderController.update_order_status(order_id, new_status)
        if success:
            pass  # clients poll /api/orders/status and /api/admin/orders/poll
        return jsonify({'success': success})

    # ── Admin AJAX: Update payment status ──────

    @app.route('/admin/order/<int:order_id>/payment', methods=['POST'])
    @admin_required
    def admin_update_payment_status(order_id):
        data = request.get_json()
        new_status = data.get('payment_status')
        if new_status not in ['Paid', 'Unpaid']:
            return jsonify({'success': False, 'message': 'Invalid payment status'})
        order = Order.query.get(order_id)
        if not order:
            return jsonify({'success': False, 'message': 'Order not found'})
        order.payment_status = new_status
        db.session.commit()
        return jsonify({'success': True})

    # ── User: Submit order rating ────────────
    @app.route('/order/<int:order_id>/rate', methods=['POST'])
    @login_required
    def submit_order_rating(order_id):
        order = Order.query.get_or_404(order_id)
        # Only the owner can rate and only after delivery
        if order.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Not authorized'})
        if order.status != 'Delivered':
            return jsonify({'success': False, 'message': 'You can only rate delivered orders'})
        if order.rating:
            return jsonify({'success': False, 'message': 'You have already rated this order'})
        data = request.get_json()
        stars = int(data.get('stars', 0))
        if stars < 1 or stars > 5:
            return jsonify({'success': False, 'message': 'Rating must be between 1 and 5 stars'})
        comment = (data.get('comment') or '').strip()[:500]
        rating = OrderRating(order_id=order_id, user_id=current_user.id, stars=stars, comment=comment or None)
        db.session.add(rating)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Thank you for your rating!'})

    # ── Admin: View all ratings ─────────────
    @app.route('/admin/ratings')
    @admin_required
    def admin_ratings():
        ratings = OrderRating.query.order_by(OrderRating.created_at.desc()).all()
        avg = round(sum(r.stars for r in ratings) / len(ratings), 1) if ratings else 0
        return render_template('admin/ratings.html', ratings=ratings, avg=avg, total=len(ratings))

    # ── Admin AJAX: Get user orders ─────────

    @app.route('/admin/user/<int:user_id>/orders')
    @admin_required
    def admin_user_orders(user_id):
        user_orders = Order.query.filter_by(user_id=user_id).order_by(Order.order_date.desc()).all()
        data = []
        for o in user_orders:
            data.append({
                'order_number': o.order_number,
                'status': o.status,
                'total_amount': o.total_amount,
                'order_date': o.order_date.strftime('%b %d, %Y %I:%M %p'),
                'items': [{'product_name': i.product_name, 'quantity': i.quantity} for i in o.items]
            })
        return jsonify({'orders': data})

    # ── Admin AJAX: Upload product image (base64 stored in DB) ────

    @app.route('/admin/product/upload-image', methods=['POST'])
    @admin_required
    def admin_upload_product_image():
        import base64
        product_id = request.form.get('product_id')
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'success': False, 'message': 'Product not found'})

        file = request.files.get('image')
        if not file or not file.filename:
            return jsonify({'success': False, 'message': 'No file uploaded'})
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': 'Invalid file type'})

        ext = file.filename.rsplit('.', 1)[1].lower()
        mime = f'image/{ext}'
        img_bytes = file.read()
        if len(img_bytes) > 5 * 1024 * 1024:
            return jsonify({'success': False, 'message': 'Image too large (max 5MB)'})

        b64 = base64.b64encode(img_bytes).decode('utf-8')
        data_url = f"data:{mime};base64,{b64}"

        product.image_data = data_url
        product.image_url = f"_data_"   # sentinel: use image_data
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': f'Save failed: {str(e)}'})

        return jsonify({'success': True, 'image_url': data_url})

    # ── Admin AJAX: Toggle product availability ─

    @app.route('/admin/product/<int:product_id>/toggle', methods=['POST'])
    @admin_required
    def admin_toggle_product(product_id):
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'success': False})
        data = request.get_json()
        product.is_available = data.get('available', True)
        db.session.commit()
        return jsonify({'success': True})

    # ── Admin: Add new product ───────────────

    @app.route('/admin/product/add', methods=['POST'])
    @admin_required
    def admin_add_product():
        import base64
        name = request.form.get('name', '').strip()
        price = request.form.get('price')
        flavor = request.form.get('flavor', '').strip()
        size = request.form.get('size', '').strip()
        description = request.form.get('description', '').strip()
        category = request.form.get('category', 'Fries').strip()

        if not all([name, price, flavor, size, description]):
            missing = [k for k,v in {'name':name,'price':price,'flavor':flavor,'size':size,'description':description}.items() if not v]
            return jsonify({'success': False, 'message': f'Missing fields: {", ".join(missing)}'})

        try:
            price = float(price)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Invalid price'})

        image_url = 'fries/cheese.png'
        image_data = None

        file = request.files.get('image')
        if file and file.filename and allowed_file(file.filename):
            ext = file.filename.rsplit('.', 1)[1].lower()
            mime = f'image/{ext}'
            img_bytes = file.read()
            if len(img_bytes) <= 5 * 1024 * 1024:
                b64 = base64.b64encode(img_bytes).decode('utf-8')
                image_data = f"data:{mime};base64,{b64}"
                image_url = '_data_'

        product = Product(
            name=name, price=price, flavor=flavor, size=size,
            description=description, image_url=image_url,
            image_data=image_data, is_available=True, category=category
        )
        db.session.add(product)
        db.session.commit()

        return jsonify({'success': True, 'product_id': product.id})

    # ── Admin: Edit product ──────────────────
    @app.route('/admin/product/<int:product_id>/edit', methods=['POST'])
    @admin_required
    def admin_edit_product(product_id):
        import base64
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'success': False, 'message': 'Product not found'})

        name = request.form.get('name', '').strip()
        price = request.form.get('price')
        flavor = request.form.get('flavor', '').strip()
        size = request.form.get('size', '').strip()
        description = request.form.get('description', '').strip()
        category = request.form.get('category', 'Fries').strip()

        if not all([name, price, flavor, size, description]):
            return jsonify({'success': False, 'message': 'All fields are required'})

        try:
            price = float(price)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Invalid price'})

        product.name = name
        product.price = price
        product.flavor = flavor
        product.size = size
        product.description = description
        product.category = category

        file = request.files.get('image')
        if file and file.filename and allowed_file(file.filename):
            ext = file.filename.rsplit('.', 1)[1].lower()
            mime = f'image/{ext}'
            img_bytes = file.read()
            if len(img_bytes) > 5 * 1024 * 1024:
                return jsonify({'success': False, 'message': 'Image must be under 5MB'})
            b64 = base64.b64encode(img_bytes).decode('utf-8')
            product.image_data = f"data:{mime};base64,{b64}"
            product.image_url = '_data_'

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': f'Save failed: {str(e)}'})

        image_url = product.image_data if product.image_data else (
            f"/static/images/{product.image_url}" if product.image_url and product.image_url != '_data_' else None
        )
        return jsonify({'success': True, 'product': product.to_dict(), 'image_url': image_url})

    # ── Admin: Delete product ────────────────
    @app.route('/admin/product/<int:product_id>/delete', methods=['POST'])
    @admin_required
    def admin_delete_product(product_id):
        try:
            from sqlalchemy import text
            with db.engine.connect() as conn:
                conn.execute(text("PRAGMA foreign_keys = OFF"))
                conn.execute(text("DELETE FROM order_items WHERE product_id = :pid"), {"pid": product_id})
                conn.execute(text("DELETE FROM cart_items WHERE product_id = :pid"), {"pid": product_id})
                result = conn.execute(text("DELETE FROM products WHERE id = :pid"), {"pid": product_id})
                conn.execute(text("PRAGMA foreign_keys = ON"))
                conn.commit()
            if result.rowcount == 0:
                return jsonify({'success': False, 'message': 'Product not found'})
            db.session.expire_all()
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'success': False, 'message': f'Delete failed: {str(e)}'})

    # ── Admin: Bulk Delete Products ─────────────
    @app.route('/admin/products/bulk-delete', methods=['POST'])
    @admin_required
    def admin_bulk_delete_products():
        data = request.get_json() or {}
        ids = data.get('ids', [])
        if not ids:
            return jsonify({'success': False, 'message': 'No product IDs provided'})
        try:
            int_ids = [int(i) for i in ids if str(i).isdigit()]
            if not int_ids:
                return jsonify({'success': False, 'message': 'No valid IDs'})
            from sqlalchemy import text
            placeholders = ','.join([':id' + str(i) for i in range(len(int_ids))])
            params = {'id' + str(i): v for i, v in enumerate(int_ids)}
            with db.engine.connect() as conn:
                conn.execute(text("PRAGMA foreign_keys = OFF"))
                conn.execute(text(f"DELETE FROM order_items WHERE product_id IN ({placeholders})"), params)
                conn.execute(text(f"DELETE FROM cart_items WHERE product_id IN ({placeholders})"), params)
                result = conn.execute(text(f"DELETE FROM products WHERE id IN ({placeholders})"), params)
                conn.execute(text("PRAGMA foreign_keys = ON"))
                conn.commit()
            db.session.expire_all()
            return jsonify({'success': True, 'deleted': result.rowcount})
        except Exception as e:
            return jsonify({'success': False, 'message': f'Bulk delete failed: {str(e)}'})

    # ── Admin: Sales Report ─────────────────
    @app.route('/admin/sales-report')
    @admin_required
    def admin_sales_report():
        from datetime import timedelta
        from collections import defaultdict
        import calendar as cal_mod
        period = request.args.get('period', 'month')
        now = ph_now()

        if period == 'today':
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            label = 'Today'
        elif period == 'week':
            start = now - timedelta(days=now.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
            label = 'This Week'
        elif period == 'month':
            start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            label = 'This Month'
        elif period == 'year':
            start = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
            label = 'This Year'
        else:
            start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            label = 'This Month'

        period_orders = Order.query.filter(Order.order_date >= start).order_by(Order.order_date.desc()).all()
        delivered = [o for o in period_orders if o.status == 'Delivered']
        total_income = sum(o.total_amount for o in delivered)
        total_orders = len(period_orders)
        total_delivered = len(delivered)
        cancelled = sum(1 for o in period_orders if o.status == 'Cancelled')

        # Auto-summary snapshots (always computed regardless of period filter)
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        year_start  = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        all_delivered = Order.query.filter(Order.status == 'Delivered').all()
        week_start  = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
        auto_today = sum(o.total_amount for o in all_delivered if o.order_date >= today_start)
        auto_week  = sum(o.total_amount for o in all_delivered if o.order_date >= week_start)
        auto_month = sum(o.total_amount for o in all_delivered if o.order_date >= month_start)
        auto_year  = sum(o.total_amount for o in all_delivered if o.order_date >= year_start)
        auto_total = sum(o.total_amount for o in all_delivered)

        # Chart data based on selected period
        if period == 'today':
            income_by  = defaultdict(float)
            orders_by  = defaultdict(int)
            for o in delivered:
                income_by[o.order_date.hour] += o.total_amount
            for o in period_orders:
                orders_by[o.order_date.hour] += 1
            chart_labels      = [f'{h:02d}:00' for h in range(24)]
            chart_data        = [income_by.get(h, 0) for h in range(24)]
            orders_chart_data = [orders_by.get(h, 0) for h in range(24)]
        elif period == 'week':
            income_by  = defaultdict(float)
            orders_by  = defaultdict(int)
            for o in delivered:
                income_by[o.order_date.weekday()] += o.total_amount
            for o in period_orders:
                orders_by[o.order_date.weekday()] += 1
            chart_labels      = ['Mon','Tue','Wed','Thu','Fri','Sat','Sun']
            chart_data        = [income_by.get(i, 0) for i in range(7)]
            orders_chart_data = [orders_by.get(i, 0) for i in range(7)]
        elif period == 'month':
            days_in = cal_mod.monthrange(now.year, now.month)[1]
            income_by  = defaultdict(float)
            orders_by  = defaultdict(int)
            for o in delivered:
                income_by[o.order_date.day] += o.total_amount
            for o in period_orders:
                orders_by[o.order_date.day] += 1
            chart_labels      = [str(d) for d in range(1, days_in + 1)]
            chart_data        = [income_by.get(d, 0) for d in range(1, days_in + 1)]
            orders_chart_data = [orders_by.get(d, 0) for d in range(1, days_in + 1)]
        else:  # year
            income_by  = defaultdict(float)
            orders_by  = defaultdict(int)
            for o in delivered:
                income_by[o.order_date.month] += o.total_amount
            for o in period_orders:
                orders_by[o.order_date.month] += 1
            chart_labels      = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
            chart_data        = [income_by.get(m, 0) for m in range(1, 13)]
            orders_chart_data = [orders_by.get(m, 0) for m in range(1, 13)]

        # Top products in period
        product_sales = defaultdict(lambda: {'qty': 0, 'revenue': 0.0})
        for o in delivered:
            for item in o.items:
                product_sales[item.product_name]['qty'] += item.quantity
                product_sales[item.product_name]['revenue'] += item.subtotal
        top_products = sorted(product_sales.items(), key=lambda x: x[1]['revenue'], reverse=True)[:10]

        pending_count = Order.query.filter(Order.status == 'Pending').count()

        return render_template('admin/sales_report.html',
                               period=period, label=label,
                               period_orders=period_orders,
                               total_income=total_income,
                               total_orders=total_orders,
                               total_delivered=total_delivered,
                               cancelled=cancelled,
                               chart_labels=chart_labels,
                               chart_data=chart_data,
                               orders_chart_data=orders_chart_data,
                               top_products=top_products,
                               pending_count=pending_count,
                               auto_today=auto_today,
                               auto_week=auto_week,
                               auto_month=auto_month,
                               auto_year=auto_year,
                               auto_total=auto_total,
                               now=now)

    # ── Admin: Export Reports ────────────────

    @app.route('/admin/report/export')
    @admin_required
    def admin_export_report():
        import csv, io
        from flask import make_response, Response
        report_type = request.args.get('type', 'orders')  # orders | products | users
        fmt = request.args.get('format', 'csv')           # csv | pdf | docx

        # ── Gather data ──────────────────────────────
        if report_type == 'products':
            rows = Product.query.order_by(Product.flavor, Product.size).all()
            headers = ['ID', 'Name', 'Category', 'Flavor', 'Size', 'Price (₱)', 'Available']
            data = [[p.id, p.name, p.category, p.flavor, p.size,
                     f'{p.price:.2f}', 'Yes' if p.is_available else 'No'] for p in rows]
            title = 'Products Report'
        elif report_type == 'users':
            rows = User.query.order_by(User.created_at.desc()).all()
            headers = ['ID', 'Username', 'Full Name', 'Email', 'Phone', 'Admin', 'Joined']
            data = [[u.id, u.username, u.full_name, u.email, u.phone or '',
                     'Yes' if u.is_admin else 'No',
                     u.created_at.strftime('%Y-%m-%d')] for u in rows]
            title = 'Users Report'
        else:  # orders
            rows = Order.query.order_by(Order.order_date.desc()).all()
            headers = ['Order #', 'Customer', 'Email', 'Phone', 'Total (₱)', 'Status', 'Payment', 'Date']
            data = [[o.order_number, o.customer_name, o.customer_email, o.customer_phone,
                     f'{o.total_amount:.2f}', o.status, o.payment_status,
                     o.order_date.strftime('%Y-%m-%d %H:%M')] for o in rows]
            title = 'Orders Report'

        generated_at = ph_now().strftime('%Y-%m-%d %H:%M')

        # ── CSV ──────────────────────────────────────
        if fmt == 'csv':
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow([title, f'Generated: {generated_at}'])
            writer.writerow([])
            writer.writerow(headers)
            writer.writerows(data)
            response = make_response(output.getvalue())
            response.headers['Content-Type'] = 'text/csv'
            response.headers['Content-Disposition'] = f'attachment; filename=potato_corner_{report_type}_{ph_now().strftime("%Y%m%d")}.csv'
            return response

        # ── PDF ──────────────────────────────────────
        elif fmt == 'pdf':
            try:
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.lib import colors
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_CENTER, TA_LEFT
            except ImportError:
                return jsonify({'error': 'reportlab not installed'}), 500

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                    rightMargin=1.5*cm, leftMargin=1.5*cm,
                                    topMargin=2*cm, bottomMargin=1.5*cm)
            styles = getSampleStyleSheet()
            brand_color = colors.HexColor('#f59e0b')
            dark_color  = colors.HexColor('#1e293b')

            title_style = ParagraphStyle('TitleStyle', parent=styles['Title'],
                                         fontSize=20, textColor=dark_color,
                                         spaceAfter=4, alignment=TA_CENTER)
            sub_style   = ParagraphStyle('SubStyle', parent=styles['Normal'],
                                         fontSize=9, textColor=colors.HexColor('#64748b'),
                                         spaceAfter=12, alignment=TA_CENTER)

            table_data = [headers] + data
            col_count  = len(headers)
            page_w     = landscape(A4)[0] - 3*cm
            col_w      = [page_w / col_count] * col_count

            tbl = Table(table_data, colWidths=col_w, repeatRows=1)
            tbl.setStyle(TableStyle([
                ('BACKGROUND',   (0, 0), (-1, 0), brand_color),
                ('TEXTCOLOR',    (0, 0), (-1, 0), colors.white),
                ('FONTNAME',     (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',     (0, 0), (-1, 0), 10),
                ('ALIGN',        (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#fef9ec')]),
                ('FONTNAME',     (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE',     (0, 1), (-1, -1), 9),
                ('GRID',         (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
                ('TOPPADDING',   (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING',(0, 0), (-1, -1), 6),
                ('LEFTPADDING',  (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('ROUNDEDCORNERS', [4]),
            ]))

            story = [
                Paragraph('🍟 Potato Corner', title_style),
                Paragraph(f'{title}  •  Generated: {generated_at}', sub_style),
                tbl,
            ]
            doc.build(story)
            buf.seek(0)
            response = make_response(buf.read())
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename=potato_corner_{report_type}_{ph_now().strftime("%Y%m%d")}.pdf'
            return response

        # ── DOCX ─────────────────────────────────────
        elif fmt == 'docx':
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Cm, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_ALIGN_VERTICAL
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
            except ImportError:
                return jsonify({'error': 'python-docx not installed'}), 500

            doc = Document()

            # Page margins
            for section in doc.sections:
                section.top_margin    = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin   = Cm(2.5)
                section.right_margin  = Cm(2.5)

            # Title
            h = doc.add_heading('🍟 Potato Corner', 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

            sub = doc.add_paragraph(f'{title}  |  Generated: {generated_at}')
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(10)
            sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            doc.add_paragraph()

            # Table
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = 'Table Grid'

            # Header row
            hdr_cells = tbl.rows[0].cells
            for i, h_txt in enumerate(headers):
                cell = hdr_cells[i]
                cell.text = h_txt
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.font.size = Pt(10)
                # Yellow background
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F59E0B')
                tc_pr.append(shd)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Data rows
            for row_idx, row in enumerate(data):
                cells = tbl.add_row().cells
                fill = 'FFFFFF' if row_idx % 2 == 0 else 'FEF9EC'
                for i, val in enumerate(row):
                    cells[i].text = str(val)
                    cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                    cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tc_pr = cells[i]._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), fill)
                    tc_pr.append(shd)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            response = make_response(buf.read())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename=potato_corner_{report_type}_{ph_now().strftime("%Y%m%d")}.docx'
            return response

        return jsonify({'error': 'Unsupported format'}), 400
